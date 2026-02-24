"""
Conversión de Word (.docx) a formato Excel estándar de HUs.
Extrae tablas o contenido estructurado y genera una hoja compatible con hu_analyzer.
También: consolidación de múltiples Excels en uno solo.
"""

import os
import re
from docx import Document
from docx.table import Table
import openpyxl
from openpyxl.styles import Font, Alignment

# Columnas estándar para HUs (orden y nombres esperados por hu_analyzer)
STANDARD_HEADERS = [
    "ID",
    "Titulo",
    "Descripción",
    "Criterios de aceptación",
    "Notas",
]

# Mapeo de posibles nombres en Word → columna estándar
HEADER_ALIASES = {
    "id": "ID",
    "id hu": "ID",
    "hu id": "ID",
    "hu-id": "ID",
    "no. hu": "ID",
    "no hu": "ID",
    "código": "ID",
    "codigo": "ID",
    "titulo": "Titulo",
    "título": "Titulo",
    "título ": "Titulo",
    "etapa/módulo": "Titulo",
    "etapa": "Titulo",
    "descripción": "Descripción",
    "descripcion": "Descripción",
    "descripción/objetivo": "Descripción",
    "historia de usuario": "Descripción",
    "definición funcional": "Descripción",
    "definicion funcional": "Descripción",
    "descripción corta": "Titulo",
    "criterios de aceptación": "Criterios de aceptación",
    "criterios": "Criterios de aceptación",
    "aceptación": "Criterios de aceptación",
    "notas": "Notas",
    "observaciones": "Notas",
    "comentarios": "Notas",
    "reglas de negocio": "Notas",
    "requerimientos ux/ui": "Notas",
}


# Palabras clave para encontrar columna de descripción/contenido principal en Excel común
DESC_KEYWORDS = ("descripción", "descripcion", "historia", "objetivo", "definición", "definicion", "contenido")
TITLE_KEYWORDS = ("titulo", "título", "etapa", "módulo", "modulo")
CRITERIA_KEYWORDS = ("criterios", "aceptación", "aceptacion")
NOTES_KEYWORDS = ("notas", "observaciones", "comentarios", "reglas")


def _normalize_header(h: str) -> str:
    """Normaliza nombre de columna para matching."""
    return re.sub(r"\s+", " ", str(h or "").strip().lower())


def _map_header(raw: str) -> str:
    """Mapea header de Word a columna estándar."""
    n = _normalize_header(raw)
    return HEADER_ALIASES.get(n, raw.strip() if raw else "")


def _extract_from_tables(doc: Document) -> tuple[list[str], list[dict]]:
    """
    Extrae HUs de tablas en el Word.
    Retorna (headers, list[dict]).
    """
    all_headers = []
    all_rows = []

    for table in doc.tables:
        if not table.rows:
            continue

        # Primera fila como headers
        raw_headers = [cell.text.strip() for cell in table.rows[0].cells]
        if not raw_headers or not any(raw_headers):
            continue

        # Mapear a estándar; si no hay ID explícito, usar primera columna
        headers = []
        id_col_idx = None
        for i, h in enumerate(raw_headers):
            mapped = _map_header(h)
            if mapped:
                headers.append(mapped)
                if mapped == "ID" or (not id_col_idx and i == 0):
                    id_col_idx = len(headers) - 1
            else:
                headers.append(h or f"Col_{i}")

        if id_col_idx is None:
            id_col_idx = 0

        # Filas de datos
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if not cells:
                continue

            row_dict = {}
            for i, h in enumerate(headers):
                val = cells[i] if i < len(cells) else ""
                if h:
                    row_dict[h] = val

            # Asegurar ID
            if "ID" not in row_dict or not str(row_dict.get("ID", "")).strip():
                row_dict["ID"] = str(cells[id_col_idx]).strip() if id_col_idx < len(cells) else ""

            if row_dict.get("ID"):
                all_rows.append(row_dict)
                if not all_headers:
                    all_headers = headers

    if all_headers and all_rows:
        return all_headers, all_rows

    return [], []


# Patrones para detectar secciones en párrafos (ej: "Titulo:", "Descripción:", "Criterios:")
SECTION_PATTERN = re.compile(
    r"^(titulo|título|descripción|descripcion|criterios|notas|observaciones|definición|definicion)\s*:?\s*(.*)$",
    re.IGNORECASE
)


def _extract_from_paragraphs(doc: Document, initiative_name: str) -> tuple[list[str], list[dict]]:
    """
    Extrae HUs de párrafos cuando no hay tablas.
    Busca patrones como "HU-001", "Historia 1", o secciones por títulos.
    Detecta también "Titulo:", "Descripción:", "Criterios:" para estructurar el contenido.
    """
    headers = ["ID", "Titulo", "Descripción", "Criterios de aceptación", "Notas"]
    hus = []
    current = {}
    current_id = ""
    buffer = []

    def _map_section_key(match_key: str) -> str:
        k = match_key.lower().strip()
        if k in ("titulo", "título"):
            return "Titulo"
        if k in ("descripción", "descripcion"):
            return "Descripción"
        if k in ("criterios", "criterios de aceptación"):
            return "Criterios de aceptación"
        if k in ("notas", "observaciones"):
            return "Notas"
        if k in ("definición", "definicion"):
            return "Descripción"
        return "Descripción"

    def flush_hu():
        nonlocal current, current_id, buffer
        if current_id or buffer:
            if not current_id:
                current_id = f"HU-{len(hus)+1}"
            desc = current.get("Descripción", "")
            if buffer:
                buffer_txt = "\n".join(buffer).strip()
                desc = (desc + "\n\n" + buffer_txt) if desc else buffer_txt
            hus.append({
                "ID": current_id,
                "Titulo": current.get("Titulo", current_id),
                "Descripción": desc or "",
                "Criterios de aceptación": current.get("Criterios de aceptación", ""),
                "Notas": current.get("Notas", ""),
            })
        current = {}
        current_id = ""
        buffer = []

    # Patrones para detectar inicio de nueva HU
    hu_pattern = re.compile(
        r"^(?:HU[- ]?)?(\d+)|^Historia\s+(?:de\s+usuario\s+)?(\d+)|^(\d+)\.\s",
        re.IGNORECASE
    )

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        match = hu_pattern.match(text)
        if match:
            flush_hu()
            g = match.groups()
            current_id = next((x for x in g if x), f"HU-{len(hus)+1}")
            if text != current_id:
                buffer = [text[len(match.group(0)):].strip()] if len(match.group(0)) < len(text) else []
            continue

        section_match = SECTION_PATTERN.match(text)
        if section_match:
            if current_id:
                key = _map_section_key(section_match.group(1))
                val = section_match.group(2).strip()
                if val:
                    current[key] = current.get(key, "") + ("\n" + val if current.get(key) else val)
            else:
                current_id = f"HU-1"
                key = _map_section_key(section_match.group(1))
                val = section_match.group(2).strip()
                if val:
                    current[key] = val
            continue

        if current_id:
            buffer.append(text)
        else:
            current_id = f"HU-1"
            buffer = [text]

    flush_hu()

    if not hus:
        # Fallback: todo el documento como una HU
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if full_text:
            hus = [{
                "ID": "HU-1",
                "Titulo": initiative_name or "Iniciativa desde Word",
                "Descripción": full_text[:2000],
                "Criterios de aceptación": "",
                "Notas": "",
            }]

    return headers, hus


def _find_best_common_column(word_key: str, common_headers: list[str]) -> str | None:
    """
    Encuentra la columna común que mejor coincide con la clave del Word.
    Usa matching por palabras clave para mapear ID, Titulo, Descripción, etc.
    """
    w = _normalize_header(word_key)
    if not w:
        return None
    # Match exacto o por alias
    for common in common_headers:
        c = _normalize_header(common)
        if not c:
            continue
        if w == c or w in c or c in w:
            return common
    # Match por palabras clave
    if any(k in w for k in ("id", "hu", "código", "codigo", "no.")):
        for h in common_headers:
            if any(k in _normalize_header(h) for k in ("id", "hu", "código", "codigo", "no.")):
                return h
    if any(k in w for k in TITLE_KEYWORDS):
        for h in common_headers:
            if any(k in _normalize_header(h) for k in TITLE_KEYWORDS):
                return h
    if any(k in w for k in DESC_KEYWORDS):
        for h in common_headers:
            if any(k in _normalize_header(h) for k in DESC_KEYWORDS):
                return h
    if any(k in w for k in CRITERIA_KEYWORDS):
        for h in common_headers:
            if any(k in _normalize_header(h) for k in CRITERIA_KEYWORDS):
                return h
    if any(k in w for k in NOTES_KEYWORDS):
        for h in common_headers:
            if any(k in _normalize_header(h) for k in NOTES_KEYWORDS):
                return h
    return None


def _map_word_hu_to_common_format(word_row: dict, common_headers: list[str], word_headers: list[str]) -> dict:
    """
    Mapea una HU del Word al formato común del Excel.
    Preserva toda la información: contenido sin columna directa va a Descripción o columna más apropiada.
    """
    result = {h: "" for h in common_headers}
    unmapped_content = []

    for wh, val in word_row.items():
        if not val or not str(val).strip():
            continue
        val_str = str(val).strip()
        common_col = _find_best_common_column(wh, common_headers)
        if common_col:
            existing = result.get(common_col, "")
            if existing:
                result[common_col] = existing + "\n\n" + val_str
            else:
                result[common_col] = val_str
        else:
            unmapped_content.append(f"[{wh}]\n{val_str}")

    if unmapped_content:
        desc_col = None
        for h in common_headers:
            if any(k in _normalize_header(h) for k in DESC_KEYWORDS):
                desc_col = h
                break
        if not desc_col:
            desc_col = common_headers[2] if len(common_headers) > 2 else common_headers[-1]
        extra = "\n\n---\n".join(unmapped_content)
        existing = result.get(desc_col, "")
        result[desc_col] = (existing + "\n\n" + extra) if existing else extra

    return result


def word_to_hus(docx_path: str, initiative_name: str = None) -> tuple[list[str], list[dict]]:
    """
    Convierte un .docx a estructura de HUs (headers, list[dict]).
    initiative_name: nombre para la hoja/iniciativa (por defecto: nombre del archivo).
    """
    if not initiative_name:
        import os
        initiative_name = os.path.splitext(os.path.basename(docx_path))[0]

    doc = Document(docx_path)
    headers, rows = _extract_from_tables(doc)

    if not rows and doc.paragraphs:
        headers, rows = _extract_from_paragraphs(doc, initiative_name)

    # Unificar headers: estándar + cualquier extra del Word
    seen = set()
    final_headers = []
    for h in STANDARD_HEADERS:
        if h not in seen:
            final_headers.append(h)
            seen.add(h)
    for h in headers:
        if h and h not in seen:
            final_headers.append(h)
            seen.add(h)

    # Rellenar columnas faltantes en cada fila
    for row in rows:
        for h in final_headers:
            if h not in row:
                row[h] = ""

    return final_headers, rows


def create_excel_sheet_from_word(
    wb: openpyxl.Workbook,
    sheet_name: str,
    headers: list[str],
    rows: list[dict],
) -> None:
    """
    Crea una hoja en el workbook con el formato esperado por hu_analyzer:
    - Filas 1-7: reservadas (metadata opcional)
    - Fila 8: headers
    - Fila 9+: datos
    """
    from hu_analyzer import HEADER_ROW, DATA_START_ROW

    if sheet_name in wb.sheetnames:
        del wb[sheet_name]
    ws = wb.create_sheet(sheet_name)

    # Fila 8: headers
    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=HEADER_ROW, column=col_idx, value=h)
        cell.font = Font(bold=True, name="Arial", size=10)
        cell.alignment = Alignment(wrap_text=True, vertical="center")

    # Filas de datos
    for row_idx, row in enumerate(rows, DATA_START_ROW):
        for col_idx, h in enumerate(headers, 1):
            val = row.get(h, "")
            ws.cell(row=row_idx, column=col_idx, value=str(val) if val else "")
        ws.row_dimensions[row_idx].height = 80

    # Ajustar anchos de columna
    from openpyxl.utils import get_column_letter
    for col_idx in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = 25


def add_word_as_sheet_to_excel(
    excel_path: str,
    docx_path: str,
    initiative_name: str = None,
    common_headers: list[str] = None,
) -> str:
    """
    Agrega el contenido del Word como nueva hoja al Excel existente.
    Si common_headers se proporciona (encabezados del Excel base), mapea el contenido
    del Word al formato común para que coincida con las otras iniciativas.
    Si el Excel no existe, lo crea. Retorna la ruta del Excel.
    """
    import os

    word_headers, word_rows = word_to_hus(docx_path, initiative_name)
    if not word_rows:
        raise ValueError("No se encontraron HUs en el documento Word.")

    if common_headers and os.path.exists(excel_path):
        rows = [_map_word_hu_to_common_format(r, common_headers, word_headers) for r in word_rows]
        headers = common_headers
    else:
        headers = word_headers
        rows = word_rows

    name = initiative_name or os.path.splitext(os.path.basename(docx_path))[0]
    sheet_name = re.sub(r'[\\/*?:\[\]]', '_', name)[:31]

    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
    else:
        wb = openpyxl.Workbook()

    create_excel_sheet_from_word(wb, sheet_name, headers, rows)

    # Si es workbook nuevo, quitar hoja por defecto vacía
    if not os.path.exists(excel_path) and len(wb.worksheets) > 1:
        for ws in list(wb.worksheets):
            if ws.title != sheet_name:
                wb.remove(ws)
                break
    wb.save(excel_path)
    return excel_path


def _copy_sheet(source_ws, target_wb: openpyxl.Workbook, sheet_name: str) -> None:
    """Copia una hoja de un workbook a otro (valores y dimensiones)."""
    if sheet_name in target_wb.sheetnames:
        return  # ya existe, no sobrescribir
    new_ws = target_wb.create_sheet(title=sheet_name)
    for row in source_ws.iter_rows():
        for cell in row:
            new_ws.cell(row=cell.row, column=cell.column, value=cell.value)
    for row_num, dim in source_ws.row_dimensions.items():
        if dim.height:
            new_ws.row_dimensions[row_num].height = dim.height
    for col_letter, dim in source_ws.column_dimensions.items():
        if dim.width:
            new_ws.column_dimensions[col_letter].width = dim.width


def merge_excel_files(excel_paths: list[str], output_path: str) -> str:
    """
    Consolida múltiples archivos Excel en uno solo.
    Cada hoja de cada archivo se agrega como pestaña. Si hay conflicto de nombres,
    se usa el nombre del archivo como prefijo (ej: "Archivo1_Iniciativa").
    Retorna la ruta del Excel consolidado.
    """
    if not excel_paths:
        raise ValueError("Se requiere al menos un archivo Excel.")
    wb_target = openpyxl.load_workbook(excel_paths[0])
    used_names = set(wb_target.sheetnames)

    for path in excel_paths[1:]:
        wb_src = openpyxl.load_workbook(path)
        base_name = os.path.splitext(os.path.basename(path))[0]
        for ws in wb_src.worksheets:
            name = ws.title
            if name in used_names:
                name = f"{base_name}_{name}"[:31]
            if name in used_names:
                idx = 1
                while f"{name}_{idx}"[:31] in used_names:
                    idx += 1
                name = f"{name}_{idx}"[:31]
            used_names.add(name)
            _copy_sheet(ws, wb_target, name)
        wb_src.close()

    wb_target.save(output_path)
    return output_path


def word_to_excel_file(docx_path: str, output_xlsx_path: str, initiative_name: str = None) -> str:
    """
    Convierte un .docx a archivo Excel con una hoja.
    Retorna la ruta del Excel generado.
    """
    headers, rows = word_to_hus(docx_path, initiative_name)
    if not rows:
        raise ValueError("No se encontraron HUs en el documento Word.")

    wb = openpyxl.Workbook()
    name = initiative_name or __import__("os").path.splitext(__import__("os").path.basename(docx_path))[0]
    # Sanitizar nombre de hoja (máx 31 chars, sin caracteres prohibidos)
    sheet_name = re.sub(r'[\\/*?:\[\]]', '_', name)[:31]
    create_excel_sheet_from_word(wb, sheet_name, headers, rows)
    # Eliminar hoja por defecto (Sheet)
    for ws in list(wb.worksheets):
        if ws.title != sheet_name:
            wb.remove(ws)
            break
    wb.save(output_xlsx_path)
    return output_xlsx_path
